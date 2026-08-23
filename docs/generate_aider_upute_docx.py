# -*- coding: utf-8 -*-
"""Generira Word upute: Aider (lokalno) + Oriphiel AI (ai.oriph.io)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm

OUT = Path(r"C:\GIT_PROJEKTI\oriphiel-platform\docs\Aider_Oriphiel_AI_upute.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
h("Aider + Oriphiel AI — lokalna instalacija i povezivanje", 0)
p(
    "Aider je lokalni AI coding agent (terminal). Oriphiel AI (Open WebUI + Ollama "
    "na VPS-u) daje modele preko OpenAI-kompatibilnog API-ja.",
    italic=True,
)

bullets([
    "Aider (lokalno na Windowsu): instalira se na PC — trenutno aider 0.86.2",
    "Oriphiel AI (udaljeno): https://ai.oriph.io",
    "VPS: root@186.240.157.80 (KVM 8)",
    "Lokalni paket Oriphiel: C:\\GIT_PROJEKTI\\oriphiel-platform",
    "Službene Aider docs: https://aider.chat/docs/",
    "Ažurirano: 9. 8. 2026.",
])

h("Sadržaj", 1)
bullets([
    "1. Što je što",
    "2. Preduvjeti",
    "3. Instalacija Aidera (Windows PowerShell)",
    "4. PATH — ako aider nije prepoznat",
    "5. API key u Oriphiel AI (Open WebUI)",
    "6. Povezivanje Aider → Oriphiel AI",
    "7. Trajna konfiguracija (.aider.conf.yml)",
    "8. Svakodnevni rad",
    "9. Modeli",
    "10. Troubleshooting",
])

# ---------------------------------------------------------------------------
h("1. Što je što", 1)
bullets([
    "Aider — lokalni program u PowerShellu; uređuje kod u projektu, predlaže diffove, radi s gitom",
    "Oriphiel AI — chat sučelje na https://ai.oriph.io (Open WebUI) + Ollama modeli na VPS-u",
    "Veza — Aider šalje upite na Open WebUI API (https://ai.oriph.io/api), ne treba lokalni Ollama",
])
p(
    "Napomena: Aider ≠ Taracod „Aiden“ (to je drugi proizvod). Ovdje se radi o aider.chat.",
    italic=True,
)

# ---------------------------------------------------------------------------
h("2. Preduvjeti", 1)
bullets([
    "Windows 10/11",
    "PowerShell 7 (preporučeno)",
    "Python 3.9–3.12 (provjereno: Python 3.12)",
    "Git (preporučeno za Aider workflow)",
    "Račun na https://ai.oriph.io (admin ili korisnik s API key-om)",
    "Oriphiel AI VPS mora biti up (Open WebUI + Ollama + modeli)",
])

# ---------------------------------------------------------------------------
h("3. Instalacija Aidera (Windows PowerShell)", 1)
p("Pokreni PowerShell na Windowsu (ne unutar SSH sesije na VPS).")

code(
    "python -m pip install aider-install\n"
    "aider-install"
)

p("Očekivano: instalira se aider-chat (npr. 0.86.2) i executable aider.")
p("Provjera verzije:")
code("aider --version")
p("Očekivano: aider 0.86.2 (ili novija).")

p("Alternativa (uv):")
code(
    "python -m pip install uv\n"
    "uv tool install --force --python python3.12 --with pip aider-chat@latest"
)

# ---------------------------------------------------------------------------
h("4. PATH — ako aider nije prepoznat", 1)
p(
    "Nakon aider-install često vidiš: "
    "warning: C:\\Users\\vittv\\.local\\bin is not on your PATH."
)

h("4.1 Trenutna PowerShell sesija", 2)
code('$env:PATH = "C:\\Users\\vittv\\.local\\bin;$env:PATH"\naider --version')

h("4.2 Novi shell", 2)
p("Zatvori PowerShell i otvori novi — aider-install ponekad već ažurira User PATH.")

h("4.3 Trajno (ako i dalje ne radi)", 2)
code(
    '[Environment]::SetEnvironmentVariable(\n'
    '  "Path",\n'
    '  [Environment]::GetEnvironmentVariable("Path","User") + ";C:\\Users\\vittv\\.local\\bin",\n'
    '  "User"\n'
    ")\n"
    "# Zatim NOVI PowerShell"
)

p("Za drugi Windows korisnički račun zamijeni vittv svojim korisničkim imenom.")

# ---------------------------------------------------------------------------
h("5. API key u Oriphiel AI (Open WebUI)", 1)
numbered([
    "Otvori https://ai.oriph.io i prijavi se",
    "Settings (Postavke) → Account (Račun) → API keys",
    "Create new secret key / Generiraj API key",
    "Kopiraj key (prikazuje se jednom — spremi ga sigurno)",
])
p(
    "Ne commitaj API key u git. Ne dijeli ga u chatu ako nije nužno.",
    italic=True,
)

# ---------------------------------------------------------------------------
h("6. Povezivanje Aider → Oriphiel AI", 1)
p("U PowerShellu (u folderu projekta na kojem želiš raditi):")

code(
    '$env:OPENAI_API_BASE = "https://ai.oriph.io/api"\n'
    '$env:OPENAI_API_KEY = "sk-TVOJ-OPEN-WEBUI-API-KEY"\n'
    "\n"
    r"cd C:\putanja\do\projekta" + "\n"
    "aider --model openai/qwen2.5-coder:14b"
)

p("Važno:", bold=True)
bullets([
    "OPENAI_API_BASE mora biti https://ai.oriph.io/api (s /api na kraju)",
    "Ime modela prefiksiraj s openai/ da litellm koristi OpenAI-kompatibilni endpoint",
    "Naziv modela mora točno odgovarati onome što Open WebUI / Ollama nudi (npr. qwen2.5-coder:14b)",
])

h("6.1 Provjera API-ja (opcionalno)", 2)
code(
    '$env:OPENAI_API_KEY = "sk-TVOJ-KEY"\n'
    "curl -H \"Authorization: Bearer $env:OPENAI_API_KEY\" https://ai.oriph.io/api/models"
)

p("Ako vidiš JSON s listom modela — veza radi.")

# ---------------------------------------------------------------------------
h("7. Trajna konfiguracija (.aider.conf.yml)", 1)
p(
    "Da ne tipkaš env varijable svaki put, u korijenu projekta (ili u kućnom direktoriju) "
    "napravi datoteku .aider.conf.yml:"
)

code(
    "openai-api-base: https://ai.oriph.io/api\n"
    "openai-api-key: sk-TVOJ-OPEN-WEBUI-API-KEY\n"
    "model: openai/qwen2.5-coder:14b"
)

p("Zatim u projektu:")
code("cd C:\\putanja\\do\\projekta\naider")

p(
    "Ako je .aider.conf.yml u projektu, dodaj ga u .gitignore ako sadrži API key "
    "(ili koristi env varijable umjesto keya u datoteci).",
    italic=True,
)

# ---------------------------------------------------------------------------
h("8. Svakodnevni rad", 1)
numbered([
    "Otvori PowerShell",
    "cd u git projekt",
    "Pokreni aider (s env ili .aider.conf.yml)",
    "Napiši zadatak na hrvatskom ili engleskom (npr. „Dodaj validaciju emaila u formu“)",
    "Pregledaj predložene izmjene — potvrdi ili odbij",
    "Koristi /help unutar Aidera za naredbe",
])

p("Korisne naredbe unutar sesije:", bold=True)
bullets([
    "/add datoteka.py — dodaj datoteku u kontekst",
    "/drop datoteka.py — makni iz konteksta",
    "/commit — commitaj izmjene",
    "/undo — poništi zadnji Aider commit",
    "/exit ili Ctrl+C — izlaz",
])

# ---------------------------------------------------------------------------
h("9. Modeli", 1)
p("Tipični modeli na Oriphiel AI (KVM 8):")
bullets([
    "openai/qwen2.5-coder:14b — kodiranje (preporučeno za Aider)",
    "openai/llama3.1:8b — opći chat / manje zahtjevni zadaci",
])
p("Promjena modela u sesiji:")
code("aider --model openai/llama3.1:8b")
p(
    "Lokalni 8B/14B modeli nisu ChatGPT — dobri su za privatnost i svakodnevni kod; "
    "za velike refaktore očekuj sporiji rad i više iteracija.",
    italic=True,
)

# ---------------------------------------------------------------------------
h("10. Troubleshooting", 1)
bullets([
    "aider nije prepoznat → §4 PATH (C:\\Users\\…\\.local\\bin)",
    "401 / Unauthorized → krivi ili istekao API key; generiraj novi u Open WebUI",
    "Connection error → provjeri https://ai.oriph.io u browseru; VPS / Nginx / Docker",
    "Model not found → točan naziv iz Open WebUI selektora; prefiks openai/",
    "Krivi base URL → mora biti …/api, ne samo https://ai.oriph.io",
    "Sporo → normalno na CPU VPS-u s 14B modelom; smanji kontekst (/drop) ili koristi 8B",
    "JSON alati u WebUI chatu ≠ Aider problem — to je Open WebUI builtin tools (vidi Oriphiel upute)",
])

h("10.1 Update Aidera", 2)
code("aider-install\n# ili\nuv tool install --force --python python3.12 --with pip aider-chat@latest")

h("10.2 Related dokumenti", 2)
bullets([
    "C:\\GIT_PROJEKTI\\oriphiel-platform\\docs\\Oriphiel_AI_OpenWebUI_upute.docx — VPS, branding, tools OFF",
    "C:\\GIT_PROJEKTI\\oriphiel-platform\\UPUTE.txt — kratki deploy podsjetnik",
    "https://aider.chat/docs/llms/openai-compat.html — OpenAI-compatible API",
])

# ---------------------------------------------------------------------------
h("Sažetak — checklist", 1)
numbered([
    "python -m pip install aider-install && aider-install",
    "Provjeri PATH → aider --version",
    "Na ai.oriph.io kreiraj API key",
    "Postavi OPENAI_API_BASE i OPENAI_API_KEY",
    "cd projekt → aider --model openai/qwen2.5-coder:14b",
])

p("Kraj dokumenta — Oriphiel d.o.o. / ai.oriph.io + Aider (lokalno)", italic=True)

doc.save(OUT)
print(OUT)
