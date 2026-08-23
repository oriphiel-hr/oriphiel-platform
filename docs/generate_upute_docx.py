# -*- coding: utf-8 -*-
"""Generira detaljne Word upute za Oriphiel AI (Ollama + Open WebUI)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"C:\GIT_PROJEKTI\oriphiel-platform\docs\Oriphiel_AI_OpenWebUI_upute.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
h("Oriphiel AI — Ollama + Open WebUI (kompletne upute)", 0)
p("Privatni AI asistent na Hostinger KVM 8. Odvojen projekt od Ravnopara (KVM 1).", italic=True)

bullets([
    "VPS hostname: srv1890026.hstgr.cloud",
    "VPS IP: 186.240.157.80",
    "Domen: https://ai.oriph.io",
    "Lokalni paket (Windows): C:\\GIT_PROJEKTI\\oriphiel-platform",
    "Firma / branding: Oriphiel d.o.o. — oriphiel.hr",
    "Ažurirano: 9. 8. 2026.",
])

h("Sadržaj", 1)
bullets([
    "1. Što je instalirano i što radi",
    "2. Lokalni paket C:\\GIT_PROJEKTI\\oriphiel-platform",
    "3. Brzi start (PowerShell deploy)",
    "4. Postavljanje VPS-a od nule (Ollama, Docker, Nginx, SSL)",
    "5. Modeli — što koristiti",
    "6. Branding (logo, CSS, top bar)",
    "7. Problem s alatima (create_tasks / write_note) i rješenje",
    "8. Korištenje Open WebUI-a (UI)",
    "9. Održavanje i troubleshooting",
    "10. Sigurnost i razdvoj od Ravnopara",
])

# ---------------------------------------------------------------------------
h("1. Što je instalirano i što radi", 1)
p("Stack na KVM 8:")
bullets([
    "DNS A: ai.oriph.io → 186.240.157.80 (bez Hostinger Website / CDN za taj subdomen)",
    "HTTPS (Let's Encrypt) preko Nginx",
    "Ollama (systemd) — lokalni LLM engine, port 11434 (nije u ufw prema internetu)",
    "Open WebUI u Dockeru (--network=host) na portu 8080",
    "Nginx proxy_pass → http://127.0.0.1:8080",
    "Branding Oriphiel (logo, custom.css, oriphiel-brand.js, top bar)",
    "Builtin tools ugašeni (builtin_tools=false, function_calling=none) da model ne šalje JSON alate",
])

p("Trenutni modeli (tipično):", bold=True)
bullets([
    "qwen2.5-coder:14b — kodiranje / tehnički rad",
    "llama3.1:8b — opći chat na hrvatskom",
])

p(
    "Važno: lokalni 8B/14B modeli nisu ChatGPT. Dobri su za privatnost i svakodnevne zadatke; "
    "za nišna znanja (npr. strip junaci) često haluciniraju ili ne znaju.",
    italic=True,
)

# ---------------------------------------------------------------------------
h("2. Lokalni paket C:\\GIT_PROJEKTI\\oriphiel-platform", 1)
p("Sve skripte i asseti za deploy žive ovdje (ne koristi stare C:\\GIT_PROJEKTI\\oriphiel-platform (zastarjelo) ni C:\\llama\\disable-tools.sh).")

code(
    r"""C:\GIT_PROJEKTI\oriphiel-platform\
  UPUTE.txt                          kratki podsjetnik
  docs\
    Oriphiel_AI_OpenWebUI_upute.docx ovaj dokument
    generate_upute_docx.py           generator Worda
  scripts\
    deploy-all.ps1                   puni deploy (branding + tools OFF)
    upload-branding.ps1              upload + apply branding
    disable-tools.ps1                samo ugasi tools
    apply-branding.sh                (Linux na VPS) branding + container
    disable-tools.sh                 (Linux na VPS) tools OFF + branding mount
  assets\
    logo.png, favicon.png, logo-white.png, logo-icon.png, logo-pattern.png
    custom.css                       UI tema (čitljivi meniji)
    oriphiel-brand.js                top bar + logo zamjena
  source\
    Logo_oriphiel.jpg                izvorni logo"""
)

p("Na VPS-u nakon deploya:")
bullets([
    "/opt/open-webui/branding/ — bind-mount u Docker",
    "/root/branding-upload/ — privremeni upload",
    "/root/oriphiel-ai/ i /root/disable-tools.sh — Linux skripte",
])

# ---------------------------------------------------------------------------
h("3. Brzi start (PowerShell na Windowsu)", 1)
p(
    "Sve PowerShell naredbe pokreći na Windowsu (ne unutar Linux SSH sesije). "
    "Pitat će root lozinku VPS-a 186.240.157.80."
)

h("3.1 Puni deploy (preporučeno)", 2)
code(
    r"powershell -ExecutionPolicy Bypass -File C:\GIT_PROJEKTI\oriphiel-platform\scripts\deploy-all.ps1"
)
p("Što radi: upload asseta + apply-branding.sh → restart Open WebUI s brandingom i tools OFF.")

h("3.2 Samo ugasi tools", 2)
code(
    r"powershell -ExecutionPolicy Bypass -File C:\GIT_PROJEKTI\oriphiel-platform\scripts\disable-tools.ps1"
)

h("3.3 Samo branding", 2)
code(
    r"powershell -ExecutionPolicy Bypass -File C:\GIT_PROJEKTI\oriphiel-platform\scripts\upload-branding.ps1"
)

h("3.4 Ručno (ekvivalent)", 2)
code(
    "scp C:\\GIT_PROJEKTI\\oriphiel-platform\\scripts\\disable-tools.sh root@186.240.157.80:/root/\n"
    "ssh root@186.240.157.80 \"sed -i 's/\\r$//' /root/disable-tools.sh; bash /root/disable-tools.sh\""
)

h("3.5 Nakon svakog deploya u pregledniku", 2)
numbered([
    "Zatvori SVE tabove ai.oriph.io",
    "Otvori novi tab: https://ai.oriph.io",
    "Ctrl+Shift+R (tvrdi refresh)",
    "Klikni Novi razgovor (ne nastavljaj stare chatove zagađene toolovima)",
    "Test: Koliko je 2+2? Odgovori samo brojem. → očekivano: 4 (običan tekst, ne JSON)",
])

# ---------------------------------------------------------------------------
h("4. Postavljanje VPS-a od nule", 1)

h("4.1 DNS", 2)
bullets([
    "hPanel → Domains → oriph.io → DNS",
    "A zapis: Name = ai , Value = 186.240.157.80",
    "Obriši ALIAS/CNAME za ai na CDN ako postoji",
])
code("nslookup ai.oriph.io")

h("4.2 SSH", 2)
code("ssh root@186.240.157.80")

h("4.3 Važno: stdout u SSH sesiji", 2)
p(
    "Ako echo, docker ps, cat ne ispisuju ništa, a nginx -t / curl greške i dalje vidiš — "
    "stdout je “ugasen”. U istoj sesiji:"
)
code("exec > /dev/tty 2>&1")
code("echo HELLO")
p(
    "Ne lijepi u terminal markdown / objašnjenja — samo čiste naredbe. "
    "Zalijepljeni komentari lome sesiju."
)

h("4.4 Update + firewall + Docker", 2)
code("apt update && apt upgrade -y")
code(
    "apt install -y curl git ufw ca-certificates docker.io docker-compose-v2 "
    "nginx certbot python3-certbot-nginx"
)
code("ufw allow OpenSSH")
code("ufw allow 80/tcp")
code("ufw allow 443/tcp")
code("ufw --force enable")
code("systemctl enable --now docker")

h("4.5 Ollama", 2)
p(
    "Ollama mora slušati na 0.0.0.0:11434 da Docker host-network može do nje. "
    "Port 11434 NE otvaraj u ufw prema internetu."
)
code("curl -fsSL https://ollama.com/install.sh | sh")
code("systemctl enable --now ollama")
code("mkdir -p /etc/systemd/system/ollama.service.d")
code(
    "cat > /etc/systemd/system/ollama.service.d/override.conf <<'EOF'\n"
    "[Service]\n"
    'Environment="OLLAMA_HOST=0.0.0.0:11434"\n'
    "EOF"
)
code("systemctl daemon-reload && systemctl restart ollama")
code("ss -lntp | grep 11434")

h("4.6 Modeli", 2)
code("ollama pull qwen2.5-coder:14b")
code("ollama pull llama3.1:8b")
code("ollama list")
p("Provjera da model radi BEZ WebUI-a:")
code('ollama run llama3.1:8b "Napiši 3 činjenice o Hrvatskoj na hrvatskom."')

h("4.7 Open WebUI (host network)", 2)
p("Koristi --network=host i port 8080. Bridge + 172.17.0.1 često ne vidi Ollamu.")
p(
    "Produkcijski način: pokreni deploy s Windowsa (deploy-all.ps1) koji diže container "
    "s brandingom i tools OFF. Ručni minimalni start (bez brandinga):"
)
code("docker rm -f open-webui")
code(
    "docker run -d --name open-webui --restart always --network=host "
    "-e OLLAMA_BASE_URL=http://127.0.0.1:11434 "
    "-e ENABLE_PERSISTENT_CONFIG=False "
    "-e 'DEFAULT_MODEL_METADATA={\"capabilities\":{\"builtin_tools\":false}}' "
    "-e 'DEFAULT_MODEL_PARAMS={\"function_calling\":\"none\"}' "
    "-v open-webui:/app/backend/data "
    "ghcr.io/open-webui/open-webui:main"
)
code("curl -I http://127.0.0.1:8080")

h("4.8 Nginx + SSL", 2)
p("sites-available/ai mora imati:")
code("proxy_pass http://127.0.0.1:8080;")
p("Ne 3000 (stari mapping).")
code("rm -f /etc/nginx/sites-enabled/default")
code("ln -sf /etc/nginx/sites-available/ai /etc/nginx/sites-enabled/ai")
code("sed -i 's/127.0.0.1:3000/127.0.0.1:8080/g' /etc/nginx/sites-available/ai")
code("nginx -t && systemctl reload nginx")
code(
    "certbot --nginx -d ai.oriph.io --email TVOJ@EMAIL.com "
    "--agree-tos --no-eff-email -v"
)
p("Ako certbot “visi” na Saving debug log — pričekaj 1–2 min, ne prekidaj odmah.")

h("4.9 Prvi login u Open WebUI", 2)
numbered([
    "Otvori https://ai.oriph.io → Get started → napravi admin račun",
    "Postavke → Povezivanja (Connections)",
    "Ollama URL: http://127.0.0.1:11434",
    "OpenAI toggle OFF ako nemaš API ključ",
    "Spremi (prozor se često ne zatvara sam — Escape / X)",
    "Odaberi model: llama3.1:8b ili qwen2.5-coder:14b",
])

# ---------------------------------------------------------------------------
h("5. Modeli — preporuke", 1)
bullets([
    "Kodiranje: qwen2.5-coder:14b — zadovoljavajuće za svakodnevni lokalni coding na KVM 8",
    "Opći chat: llama3.1:8b OK; još bolje opći qwen2.5:14b ako želiš bolje znanje (opcionalno)",
    "Ne povlači sve verzije — drži trenutni + eventualno jedan kandidat (max 2–3 modela)",
    "Llama 70B nije praktična na ovom CPU VPS-u",
    "Veći model ≠ ChatGPT kvaliteta na CPU-u; samo sporije",
])

# ---------------------------------------------------------------------------
h("6. Branding (Oriphiel)", 1)
p("Cilj: tamna tema, Oriphiel top bar, blagi watermark loga, čitljivi izbornici.")

h("6.1 Što radi branding", 2)
bullets([
    "custom.css — pozadina, top bar padding, solidne pozadine za menije/dropdown (v23)",
    "oriphiel-brand.js — ubacuje top bar (Oriphiel d.o.o. + tagline + oriphiel.hr), zamjenjuje logo",
    "logo*.png / favicon — mountani u /app/.../static/",
])

h("6.2 Tamna tema u UI", 2)
p(
    "Tema nije u Sučelje/Prilagodba nego: Postavke korisnika → Općenito → Theme → Tamna / Dark."
)

h("6.3 Problemi koje smo riješili", 2)
bullets([
    "Agresivan transparent CSS → + meni nečitljiv (tekstovi se preklapaju) → riješeno solidnim pozadinama menija",
    "Višestruki logotipi (SVG + default + Oriphiel) → JS zamjena + CSS hide",
    "Logo mount na krivi path → mount i u backend static i u build/static",
])

h("6.4 Deploy brandinga", 2)
code(
    r"powershell -ExecutionPolicy Bypass -File C:\GIT_PROJEKTI\oriphiel-platform\scripts\upload-branding.ps1"
)

# ---------------------------------------------------------------------------
h("7. Problem s alatima (Builtin Tools) — detaljno", 1)

h("7.1 Simptom", 2)
p(
    "Umjesto običnog odgovora model ispisuje JSON ili “View Result from create_tasks / "
    "write_note / replace_note_content”. Primjer: pitanje 2+2 → JSON umjesto 4."
)
p(
    "Dokaz da Ollama radi: ista pitanja u terminalu (ollama run …) daju normalan tekst. "
    "Problem je bio Open WebUI koji modelu gurа popis ugrađenih alata."
)

h("7.2 Što NIJE mjesto za isključivanje", 2)
bullets([
    "+ izbornik (Prijenos datoteka, Capture, Attach…) — to su privitci, ne create_tasks",
    "Code Interpreter OFF — samo interpreter; ne gasi note/task alate",
    "Admin → Functions (Functions 0) — custom funkcije; prazno je OK",
    "Workspace → Modeli 0 / Alati 0 — Ollama modeli nisu “Workspace Models” dok ih ne kreiraš kao zapis",
    "Admin → Općenito (Community Sharing, Memories…) — nije Capabilities Tools",
])

h("7.3 Što JEST rješenje", 2)
bullets([
    "Docker env: DEFAULT_MODEL_METADATA s builtin_tools:false",
    "Docker env: DEFAULT_MODEL_PARAMS s function_calling:none",
    "ENABLE_PERSISTENT_CONFIG=False — inače WebUI zadrži stare postavke i zanemari env",
    "Nakon toga OBAVEZNO Novi razgovor (stari chatovi ostaju “zarazeni”)",
])
p("Skripta:")
code(
    r"powershell -ExecutionPolicy Bypass -File C:\GIT_PROJEKTI\oriphiel-platform\scripts\disable-tools.ps1"
)
p(
    "Nakon uspješnog pokretanja (GOTOVO) test u novom chatu: 2+2 → 4, "
    "i pitanja o Hrvatskoj kao običan tekst — potvrđeno 9. 8. 2026."
)

h("7.4 Opcija u UI (alternativa)", 2)
numbered([
    "Radna ploča → Modeli → Create → bazni model llama3.1:8b",
    "Capabilities → ugasi Tools",
    "Spremi i u chatu biraj taj novi model",
])

# ---------------------------------------------------------------------------
h("8. Korištenje Open WebUI-a", 1)
bullets([
    "Novi razgovor — uvijek za čist kontekst",
    "Selektor modela (llama3.1:8b / qwen2.5-coder:14b) desno uz input",
    "+ — privitci (datoteke, Capture, Knowledge…)",
    "Notes / Radna ploča u sidebaru — značajke UI-a; ne znače da su toolovi u chatu aktivni",
    "Stare chatove s JSON alatima ne nastavljaj za “ozbiljan” rad",
])

# ---------------------------------------------------------------------------
h("9. Održavanje i troubleshooting", 1)

h("9.1 Korisne naredbe na VPS-u", 2)
code("docker ps -a")
code("docker logs --tail 80 open-webui")
code("docker restart open-webui")
code("systemctl status ollama --no-pager")
code("ollama list")
code("certbot certificates")
code("curl -I http://127.0.0.1:8080")
code("curl -I https://ai.oriph.io")

h("9.2 Tipični problemi", 2)
bullets([
    "JSON alati u odgovoru → disable-tools.ps1 + Novi razgovor",
    "Nečitljiv + meni → deploy branding v23+ (solidne pozadine menija)",
    "Modeli se ne vide → Connections Ollama URL http://127.0.0.1:11434 ; docker --network=host",
    "SSH bez outputa → exec > /dev/tty 2>&1 ili nova sesija",
    "scp/ssh iz Cursor agenta često ne može unijeti lozinku → pokreni PowerShell lokalno",
    "Workspace Modeli 0 → normalno za čiste Ollama modele",
])

h("9.3 Regeneracija ovog Word dokumenta", 2)
code(r"python C:\GIT_PROJEKTI\oriphiel-platform\docs\generate_upute_docx.py")
p("Izlaz: C:\\GIT_PROJEKTI\\oriphiel-platform\\docs\\Oriphiel_AI_OpenWebUI_upute.docx")

# ---------------------------------------------------------------------------
h("10. Sigurnost i razdvoj od Ravnopara", 1)
bullets([
    "Ne otvaraj 11434 u ufw prema internetu",
    "Open WebUI samo preko HTTPS + login",
    "KVM 8 (186.240.157.80) = samo AI stack",
    "KVM 1 / ravnopar.com = zaseban projekt — nema dijeljenja podataka s ovim WebUI-jem",
    "Ne commitaj root lozinke u git / chat; rotiraj ako su dijeljene",
])

h("Sažetak — dnevni workflow", 1)
numbered([
    "Otvori https://ai.oriph.io",
    "Novi razgovor",
    "Odaberi model (coder za kod, llama za chat)",
    "Pitaj normalnim jezikom",
    "Ako opet vidiš JSON alate → disable-tools.ps1 pa Novi razgovor",
    "Za UI/branding izmjene → upload-branding.ps1 ili deploy-all.ps1",
])

p("Kraj dokumenta — Oriphiel d.o.o. / ai.oriph.io", italic=True)

doc.save(OUT)
# kompatibilnost: kopija uz generator
print(OUT)
