#!/usr/bin/env python3
"""Generira Oriphiel-Messaging-Upute.docx"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Oriphiel-Messaging-Upute.docx"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def h(text, level=1):
        return doc.add_heading(text, level=level)

    def p(text, bold=False, size=11):
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold)
        para.paragraph_format.space_after = Pt(6)
        return para

    def bullet(text):
        para = doc.add_paragraph(text, style="List Bullet")
        for run in para.runs:
            set_run_font(run, size=11)
        return para

    def code_block(text):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run(text)
        set_run_font(run, size=8, name="Consolas")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3F4F6")
        shd.set(qn("w:val"), "clear")
        para._p.get_or_add_pPr().append(shd)
        return para

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = htxt
            shade_cell(cell, "1F4E79")
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                cell.text = str(val)
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, size=9)
        doc.add_paragraph()
        return table

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("Oriphiel Messaging"), size=28, bold=True, color=(31, 78, 121))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        sub.add_run("IMAP · n8n · Postgres · Ollama — detaljne upute"),
        size=14,
        color=(80, 80, 80),
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        meta.add_run(
            "Account: mario.vitt@oriphiel.hr  |  VPS: srv1890026 (186.240.157.80)  |  n8n: :5678"
        ),
        size=10,
        color=(100, 100, 100),
    )

    p(
        "Dokument opisuje arhitekturu, dijagrame toka, lokalne i VPS putanje, "
        "import u n8n, AI polja, backfill i česte greške."
    )

    # 1 Architecture
    h("1. Dijagram sustava (arhitektura)", 1)
    p(
        "Cijeli sustav ima tri sloja: Windows (razvoj/upload), VPS skripte, "
        "i runtime (IMAP, baza, disk, Ollama, n8n)."
    )
    code_block(
        """
┌─────────────────────────────────────────────────────────────────────────┐
│  WINDOWS (C:\\GIT_PROJEKTI\\oriphiel-platform\\scripts\\oriphiel_messaging\\)             │
│  • Deploy-ImapBackfill.ps1   →  scp upload na VPS                       │
│  • n8n\\*.json                →  Import u n8n UI (browser)               │
│  • ovaj .docx / UPUTE.md     →  dokumentacija                           │
└───────────────────────────────┬─────────────────────────────────────────┘
                                │ scp / import
                                ▼
┌─────────────────────────────────────────────────────────────────────────┐
│  VPS srv1890026  —  /root/oriphiel-ai/oriphiel_messaging/               │
│  • accounts/oriphiel.hr-mario.vitt.env   (lozinka + postavke)           │
│  • sync-imap-backfill.sh  →  sync-imap-backfill.py                      │
│  • ollama-enrich-message.py              (AI)                           │
│  • run-backfill-mario.vitt.sh            (jedna naredba)                │
└───────────────────────────────┬─────────────────────────────────────────┘
                                │
          ┌─────────────────────┼─────────────────────┐
          ▼                     ▼                     ▼
   ┌─────────────┐      ┌──────────────┐      ┌──────────────┐
   │ Hostinger   │      │  Postgres    │      │   Ollama     │
   │ IMAP :993   │      │  oriphiel-   │      │  :11434      │
   │ INBOX       │      │  postgres    │      │  llama3.1    │
   └─────────────┘      │  DB:oriphiel │      └──────────────┘
                        └──────┬───────┘
                               │
                        ┌──────▼───────┐
                        │ Disk         │
                        │ /var/lib/    │
                        │ oriphiel/    │
                        │ attachments/ │
                        │ email/{id}/  │
                        └──────────────┘

   n8n (http://186.240.157.80:5678)
   • Live Mail (per account): IMAP → AI → attachmenti
   • Mail Hub (jedan): backfill / enrich / count
   • Error: Oriphiel — Error notify
   • DEPRECATED: New mails - mario.vitt@….json
""".strip(
            "\n"
        )
    )

    # 2 Live
    h("2. Dijagram — Live (novi mail)", 1)
    p(
        "Workflow: Oriphiel — Live Mail (per account). "
        "Postavi Account Email + IMAP credential. Jedna kopija po mailboxu."
    )
    code_block(
        """
  Email Trigger (IMAP)
           │
           ▼
     Account Email         ← Set: ACCOUNT_EMAIL (npr. mario.vitt@…)
           │
           ▼
      Get Account          ← channels_accounts WHERE address = ACCOUNT_EMAIL
           │
           ▼
  Normalize Mail (Live)    ← from, subject, body, thread_key, external_id
           │
           ▼
    Upsert Contact         ← contacts.primary_email
           │
           ▼
    Insert Message         ← messages (subject, body, thread_key…)
           │
           ▼
  AI Enrich (Ollama)       ← HTTP → Ollama /api/generate
           │
           ▼
  Update Message AI        ← ai_summary, ai_priority, ai_draft
           │
           ▼
  If Has Attachments?
      │              │
     da             ne → kraj
      │
      ▼
  Create Attachment Directory  (/var/lib/oriphiel/attachments/email/{account_id}/)
      │
      ▼
  Prepare → Write file → Insert Attachment  (message_attachments)
""".strip(
            "\n"
        )
    )

    # 3 Backfill
    h("3. Dijagram — Backfill / AI enrich (Mail Hub)", 1)
    p(
        "Ops ide kroz Oriphiel — Mail Hub (jedan workflow). "
        "Account Config: ACTION=backfill|enrich|count + ACCOUNT_EMAIL + ENV_FILE."
    )
    code_block(
        """
  [n8n] Mail Hub — Manual Trigger
              │
              ▼
       Account Config     ← ACCOUNT_EMAIL, ENV_FILE, ACTION, RUN_AI, RESET_BEFORE
              │
              ▼
       Switch ACTION
         │         │         │
    backfill    enrich     count
         │         │         │
         ▼         ▼         ▼
      SSH nohup  SSH nohup  Postgres SELECT
   sync-imap-   enrich-     mailova / s_ai
   backfill.sh  existing-
                messages.py

  Status: /tmp/oriphiel-imap-backfill-{STATUS_TAG}.json
          /tmp/oriphiel-ai-enrich-{STATUS_TAG}.json
""".strip(
            "\n"
        )
    )

    # 4 AI
    h("4. AI polja (gdje su)", 1)
    p("Tri polja su stupci u Postgres tablici messages (baza oriphiel):")
    add_table(
        ["Stupac", "Tip", "Sadržaj"],
        [
            ["ai_summary", "TEXT", "Sažetak maila (1–3 rečenice, HR)"],
            ["ai_priority", "TEXT", "low | normal | high | urgent"],
            ["ai_draft", "TEXT", "Nacrt odgovora (2–4 rečenice)"],
        ],
    )
    p("Punjenje:")
    bullet("Live n8n: AI Enrich (Ollama) → Update Message AI")
    bullet("Backfill: ollama-enrich-message.py kad je RUN_AI=1")
    p("Pregled u bazi:")
    code_block(
        'docker exec -it oriphiel-postgres psql -U oriphiel -d oriphiel -c "\n'
        "SELECT id, left(subject,40), ai_priority, left(ai_summary,60)\n"
        "FROM messages WHERE ai_summary IS NOT NULL\n"
        'ORDER BY id DESC LIMIT 10;"'
    )

    # 5 Paths
    h("5. Putanje i datoteke", 1)
    h("5.1 Windows (lokalno)", 2)
    code_block(r"C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\\")
    add_table(
        ["Datoteka", "Uloga"],
        [
            ["sync-imap-backfill.sh / .py", "Glavni backfill"],
            ["ollama-enrich-message.py", "Ollama AI (backfill + enrich)"],
            ["enrich-existing-messages.py", "AI na već ubačenim mailovima"],
            ["run-backfill-mario.vitt.sh", "Jedna naredba: wipe + uvoz"],
            ["Deploy-ImapBackfill.ps1", "Upload skripti na VPS (s Windowsa)"],
            ["accounts\\*.env.example", "Predložak konfiguracije"],
            ["n8n\\oriphiel-live-account-stub.json", "Live IMAP+AI (po accountu)"],
            ["n8n\\oriphiel-mail-hub.json", "Hub: backfill/enrich/count"],
            ["n8n\\oriphiel-error-notify.json", "Error workflow"],
            ["n8n\\new-mails-mario-vitt.json", "DEPRECATED — ne koristi"],
            ["sql\\oriphiel-messaging-schema.sql", "Shema baze"],
            ["Check-OriphielMessaging.ps1", "Provjera baze + disk + backfill"],
            ["check-messaging.sh", "Provjera baze + disk (na VPS-u)"],
            ["Invoke-OriphielSql.ps1", "Ad-hoc SQL na oriphiel bazi"],
            ["n8n\\UPUTE.md / MULTI-ACCOUNT.md", "Dokumentacija"],
        ],
    )

    h("5.3 Provjera baze i priloga na disku", 2)
    p(
        "Koristi Check skripte — brojevi u Postgresu, AI polja, attachmenti u bazi "
        "vs datoteke na disku, i backfill status JSON."
    )
    p("S Windowsa:")
    code_block(
        "powershell -NoProfile -ExecutionPolicy Bypass -File `\n"
        r"  C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\Check-OriphielMessaging.ps1"
    )
    p("Na VPS-u:")
    code_block(
        "cd /root/oriphiel-ai/oriphiel_messaging\n"
        "bash check-messaging.sh\n"
        "# ili:\n"
        "pwsh -File ./Check-OriphielMessaging.ps1 -Local"
    )
    p("Što Check ispisuje:")
    bullet("broj accounts / contacts / messages / messages_with_ai / attachments_db")
    bullet("po accountu: messages, with_ai, attachments_db")
    bullet("zadnje poruke + ai_priority / ai_summary")
    bullet("backfill STATUS_FILE JSON + pgrep proces")
    bullet("disk: /var/lib/oriphiel/attachments/email/{account_id}/ — broj datoteka + veličina")
    bullet("usporedba: attachments_db vs files_on_disk")
    p("Ad-hoc SQL:")
    code_block(
        "powershell -ExecutionPolicy Bypass -File `\n"
        r"  C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\Invoke-OriphielSql.ps1 `\n"
        '  -Sql "SELECT count(*) FROM messages;"'
    )
    p("Svi korisni SELECT-ovi (datoteka sql/useful-selects.sql):")
    code_block(
        "# Windows (ako SSH ključ radi)\n"
        "powershell -File .\\Check-OriphielMessaging.ps1 -RunUsefulSql\n\n"
        "# VPS\n"
        "docker exec -i oriphiel-postgres psql -U oriphiel -d oriphiel \\\n"
        "  -f - < /root/oriphiel-ai/oriphiel_messaging/sql/useful-selects.sql"
    )
    p("Primjeri SELECT-ova:")
    code_block(
        "-- Brojevi\n"
        "SELECT count(*) FROM messages;\n"
        "SELECT count(*) FROM message_attachments;\n\n"
        "-- Zadnji mailovi + AI\n"
        "SELECT id, left(subject,50), ai_priority, left(ai_summary,60), received_at\n"
        "FROM messages ORDER BY id DESC LIMIT 20;\n\n"
        "-- Urgent/high\n"
        "SELECT id, subject, ai_priority FROM messages\n"
        "WHERE ai_priority IN ('urgent','high') ORDER BY received_at DESC;\n\n"
        "-- Attachmenti\n"
        "SELECT filename, size_bytes, storage_path\n"
        "FROM message_attachments ORDER BY id DESC LIMIT 20;\n\n"
        "-- Threadovi s više poruka\n"
        "SELECT thread_key, count(*) FROM messages\n"
        "WHERE thread_key IS NOT NULL GROUP BY thread_key HAVING count(*)>1\n"
        "ORDER BY 2 DESC LIMIT 20;"
    )

    h("5.2 VPS", 2)
    add_table(
        ["Što", "Putanja"],
        [
            ["Skripte", "/root/oriphiel-ai/oriphiel_messaging/"],
            ["Env (lozinka)", ".../accounts/oriphiel.hr-mario.vitt.env"],
            ["Attachmenti", "/var/lib/oriphiel/attachments/email/{account_id}/"],
            ["Status backfill", "/tmp/oriphiel-imap-backfill-mario.vitt.json"],
            ["n8n UI", "http://186.240.157.80:5678"],
            ["Ollama", "http://127.0.0.1:11434"],
            ["Postgres", "container oriphiel-postgres / DB oriphiel / user oriphiel"],
        ],
    )

    # 6 n8n
    h("6. Import JSON-ova u n8n", 1)
    p("JSON se importaju u browseru (n8n UI), ne preko SSH. Koristi samo Live + Hub (+ Error).")

    h("6.1 Oriphiel — Error notify", 2)
    bullet("Import oriphiel-error-notify.json → Active")

    h("6.2 Live Mail (per account)", 2)
    bullet(r"Import Downloads\Oriphiel-Live-Account-Stub.json")
    bullet("Account Email = mario.vitt@oriphiel.hr (ili drugi)")
    bullet("IMAP credential za taj mailbox; Postgres + SSH")
    bullet("Settings → Error Workflow → Oriphiel — Error notify")
    bullet("Active. Stari New mails - mario.vitt… → Deactivate")

    h("6.3 Mail Hub (ops — svi accounti)", 2)
    bullet(r"Import Downloads\Oriphiel-Mail-Hub.json")
    bullet("Account Config: ACTION=count|backfill|enrich + ACCOUNT_EMAIL + ENV_FILE")
    bullet("Manual Trigger")
    bullet("Novi mailbox: samo promijeni Config (ne dupliciraj Hub)")

    h("6.4 Deprecated", 2)
    p(
        "New mails - mario.vitt@oriphiel.hr.json — ne koristi. "
        "Zamijenjen s Live Stub + Mail Hub.",
        bold=True,
    )

    h("6.5 Error Workflow + sticky", 2)
    bullet("Na Live Mail workflowu: Settings → Error Workflow → Oriphiel — Error notify")
    bullet("Žuta Sticky Note je samo dokumentacija — možeš obrisati")

    # 7 Backfill howto
    h("7. Pokretanje backfilla", 1)

    h("7.1 Env datoteka (obavezno)", 2)
    p("U .env idu SAMO varijable (IME=vrijednost). NE zalijepljuj bash naredbe.")
    code_block(
        "# Oriphiel (oriphiel.hr) - mario.vitt@oriphiel.hr\n"
        "IMAP_HOST=imap.hostinger.com\n"
        "IMAP_PORT=993\n"
        "IMAP_USER=mario.vitt@oriphiel.hr\n"
        "IMAP_PASSWORD='lozinka_u_navodnicima'\n"
        "ACCOUNT_EMAIL=mario.vitt@oriphiel.hr\n"
        "RUN_AI=0\n"
        "BATCH_SIZE=50\n"
        "MARK_AS_SEEN=0\n"
        "STATUS_FILE=/tmp/oriphiel-imap-backfill-mario.vitt.json"
    )
    p("Ili u n8n Mail Hub: ACTION=backfill, RESET_BEFORE=0|1, RUN_AI=0|1.")

    h("7.2 Upload skripti s Windowsa", 2)
    p("scp s C:\\ putanjama mora ići iz Windows PowerShell-a, NE iz SSH sesije na VPS-u.")
    code_block(
        "powershell -ExecutionPolicy Bypass -File "
        r"C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\Deploy-ImapBackfill.ps1"
    )
    p("Ili pojedinačno:")
    code_block(
        r'scp "C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\sync-imap-backfill.sh" '
        "root@186.240.157.80:/root/oriphiel-ai/oriphiel_messaging/\n"
        r'scp "C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\sync-imap-backfill.py" '
        "root@186.240.157.80:/root/oriphiel-ai/oriphiel_messaging/\n"
        r'scp "C:\GIT_PROJEKTI\oriphiel-platform\scripts\oriphiel_messaging\ollama-enrich-message.py" '
        "root@186.240.157.80:/root/oriphiel-ai/oriphiel_messaging/"
    )

    h("7.3 Pokreni na VPS-u (Terminal 1)", 2)
    code_block(
        "cd /root/oriphiel-ai/oriphiel_messaging\n"
        "ENV_FILE=accounts/oriphiel.hr-mario.vitt.env \\\n"
        "RESET_BEFORE=1 \\\n"
        "RUN_AI=1 \\\n"
        "STATUS_FILE=/tmp/oriphiel-imap-backfill-mario.vitt.json \\\n"
        "bash sync-imap-backfill.sh"
    )
    p("Ili jedna naredba (nakon što je run-backfill-mario.vitt.sh uploadan):")
    code_block("bash /root/oriphiel-ai/oriphiel_messaging/run-backfill-mario.vitt.sh")

    h("7.4 Status (Terminal 2)", 2)
    p("Tek NAKON što Terminal 1 ispiše === STARTED:")
    code_block("watch -n2 cat /tmp/oriphiel-imap-backfill-mario.vitt.json")

    h("7.5 Test s limitom", 2)
    code_block(
        "LIMIT=20 DRY_RUN=1 ENV_FILE=accounts/oriphiel.hr-mario.vitt.env bash sync-imap-backfill.sh"
    )

    h("7.6 Što radi RESET_BEFORE=1", 2)
    bullet("Briše poruke i attachment zapise za taj account u bazi")
    bullet("Briše datoteke u /var/lib/oriphiel/attachments/email/{account_id}/")
    bullet("Zatim uvozi cijeli INBOX iznova")

    # 8 Mark as read
    h("8. Mark as read (MARK_AS_SEEN)", 1)
    p("Default: 0 (isključeno).", bold=True)
    p(
        "IMAP \\Seen je samo flag na mail serveru — Outlook ga prikazuje kao pročitano. "
        "To NE znači da si ti osobno pročitao mail."
    )
    bullet("Live n8n: NE stavljaj Post-process Action = Mark as read")
    bullet(
        "Backfill: MARK_AS_SEEN=0 po defaultu; uključi 1 samo ako namjerno želiš "
        "maknuti „nepročitano“ u klijentu"
    )

    # 9 Thread
    h("9. Thread grouping (thread_key)", 1)
    p("Polje messages.thread_key grupira odgovore u isti thread.")
    bullet("1. In-Reply-To (prvi Message-ID)")
    bullet("2. inače prvi ID iz References")
    bullet("3. inače Message-ID same poruke")
    bullet("4. inače uid-{broj}")

    # 10 Feature matrix
    h("10. Pregled featurea", 1)
    add_table(
        ["Feature", "Live n8n", "Python backfill", "Napomena"],
        [
            ["AI summary/priority/draft", "Da", "Da (RUN_AI=1)", "Ollama na VPS-u"],
            ["Error workflow", "Da (poveži Settings)", "—", "oriphiel-error-notify.json"],
            ["thread_key", "Da", "Da", "In-Reply-To / References"],
            ["Mark as read", "Ne (namjerno)", "Opcionalno (default 0)", "MARK_AS_SEEN"],
            ["Batch po 50", "—", "Da (BATCH_SIZE)", "Nije SplitInBatches node"],
            ["Attachmenti na disk", "Da", "Da", "/var/lib/oriphiel/attachments/..."],
        ],
    )

    # 11 Troubleshoot
    h("11. Česte greške", 1)
    add_table(
        ["Simptom", "Uzrok", "Rješenje"],
        [
            [
                "syntax error near '(' u .env",
                "Zalijepljena bash naredba ili lozinka bez navodnika",
                "Očisti .env; IMAP_PASSWORD='...'",
            ],
            [
                "Could not resolve hostname c",
                "scp C:\\... pokrenut S VPS-a",
                "scp iz Windows PowerShell-a",
            ],
            [
                "No such file status json",
                "Backfill nije krenuo",
                "Prvo bash sync-imap-backfill.sh, pa watch",
            ],
            [
                "source u sync-imap-backfill.sh",
                "Stara skripta na VPS-u",
                "Upload nove (load_env_file)",
            ],
            [
                "ai_* su NULL",
                "RUN_AI=0 ili nema Ollame / AI nodova",
                "RUN_AI=1; import novog JSON-a; Ollama up",
            ],
            [
                "Error Workflow nije u Settings",
                "Gledaš server Settings",
                "Settings UNUTAR New mails workflowa",
            ],
        ],
    )

    # 12 Checklist
    h("12. Checklist za produkciju", 1)
    bullet("Env datoteka na VPS-u ispravna (samo KEY=value, lozinka u '')")
    bullet("Nove skripte uploadane (/root/oriphiel-ai/oriphiel_messaging/)")
    bullet("grep load_env_file sync-imap-backfill.sh — postoji (nije source)")
    bullet("n8n: import New mails + Error notify; Error Workflow povezan; Active")
    bullet("Ollama radi: curl http://127.0.0.1:11434/api/tags")
    bullet("Backfill test: LIMIT=5, pa full RESET_BEFORE=1")
    bullet("Provjera: bash check-messaging.sh  (baza + disk + AI)")
    bullet("Provjera AI: SELECT ai_summary FROM messages …")

    # 13 Next
    h("13. Prijedlozi za kasnije", 1)
    bullet("Admin UI — inbox s threadovima i AI draft gumbom")
    bullet("Webhook reprocess — ponovno AI za jedan message_id")
    bullet("Slack/Telegram za ai_priority=urgent")
    bullet("Sent folder sync — outbound poruke")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        footer.add_run("\n— Kraj dokumenta — Oriphiel Messaging —"),
        size=9,
        color=(120, 120, 120),
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
