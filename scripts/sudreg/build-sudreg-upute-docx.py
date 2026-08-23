#!/usr/bin/env python3
"""Generira Sudreg-Upute.docx"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Sudreg-Upute.docx"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def h(text, level=1):
        return doc.add_heading(text, level=level)

    def p(text, bold=False, size=11):
        para = doc.add_paragraph()
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold)
        para.paragraph_format.space_after = Pt(6)
        return para

    def bullet(text):
        para = doc.add_paragraph(text, style="List Bullet")
        for run in para.runs:
            set_run_font(run, size=11)
        return para

    def code_block(text):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run(text)
        set_run_font(run, size=8, name="Consolas")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3F4F6")
        shd.set(qn("w:val"), "clear")
        para._p.get_or_add_pPr().append(shd)
        return para

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = htxt
            shade_cell(cell, "1F4E79")
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.rows[r_i + 1].cells[c_i]
                cell.text = str(val)
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, size=9)
        doc.add_paragraph()
        return table

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("Sudreg"), size=28, bold=True, color=(31, 78, 121))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        sub.add_run("Hrvatski sudski registar → Postgres — detaljne upute"),
        size=14,
        color=(80, 80, 80),
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        meta.add_run(
            "VPS: srv1890026 (186.240.157.80)  |  DB: sudreg  |  "
            "container: oriphiel-postgres  |  cron: 06:15"
        ),
        size=10,
        color=(100, 100, 100),
    )

    p(
        "Sync Sudreg API (promjene + scrape subjekata) u Postgres bazu sudreg "
        "na Ollama/Oriphiel VPS. NE na Ravnopar VPS. Baza sudreg je odvojena od "
        "oriphiel (mail/messaging)."
    )

    # 1 Architecture
    h("1. Dijagram sustava", 1)
    code_block(
        """
┌─────────────────────────────────────────────────────────────────────────┐
│  WINDOWS  C:\\GIT_PROJEKTI\\oriphiel-platform\\                                         │
│  scripts\\sudreg\\   — PowerShell skripte + Deploy-SudregToVps.ps1        │
│  data\\sudreg\\      — lokalni JSON cache, lock, logovi (ako radi lokalno)│
└───────────────────────────────┬─────────────────────────────────────────┘
                                │ Deploy / SSH (scp)
                                ▼
┌─────────────────────────────────────────────────────────────────────────┐
│  VPS srv1890026  (186.240.157.80)                                       │
│  /opt/oriphiel-ai/scripts/sudreg/   — skripte + vps-install             │
│  /opt/oriphiel-ai/data/sudreg/      — JSON, lock, progress, logovi      │
│  cron 06:15 → run-sudreg-daily.sh                                       │
└───────────────────────────────┬─────────────────────────────────────────┘
                                │
          ┌─────────────────────┼─────────────────────┐
          ▼                     ▼                     ▼
   ┌─────────────┐      ┌──────────────┐      ┌──────────────┐
   │ Sudreg API  │      │  Postgres    │      │  Disk data   │
   │ OAuth token │      │  oriphiel-   │      │  /opt/.../   │
   │ /snapshots  │      │  postgres    │      │  data/sudreg │
   │ /promjene   │      │  DB: sudreg  │      │  run/logs/   │
   │ /subjekt    │      │  user:oriphi │      └──────────────┘
   └─────────────┘      └──────────────┘
""".strip(
            "\n"
        )
    )

    # 2 Model
    h("2. Model rada", 1)

    h("2.1 Prvi put (baza prazna) — bootstrap", 2)
    code_block(
        """
  Get-SudregToken          → OAuth (cache)
         │
         ▼
  Fetch /promjene          → svi dostupni snapshotovi (JSON)
         │
         ▼
  Bootstrap baseline       → scrape SVIH MBS iz NAJSTARIJEG snapshota
         │
         ▼
  Sync delta               → scrape razlike (added MBS) do najnovijeg
         │
         ▼
  Spremi last_imported_snapshot_id  u bazi / meta
""".strip(
            "\n"
        )
    )

    h("2.2 Svaki dan (baza ima podatke) — delta", 2)
    code_block(
        """
  Fetch nove /promjene (ako treba)
         │
         ▼
  Usporedi last_imported  ↔  najnoviji snapshot
         │
         ▼
  Scrape samo DELTU (added MBS)
         │
         ▼
  UPSERT u Postgres + ažuriraj last_imported
""".strip(
            "\n"
        )
    )
    p(
        "Dok traje unos, drugi pokušaj se odbija (lock u data/sudreg/run/lock.json).",
        bold=True,
    )

    # 3 Paths
    h("3. Putanje", 1)
    add_table(
        ["Što", "Windows", "VPS"],
        [
            [
                "Skripte",
                "C:\\GIT_PROJEKTI\\oriphiel-platform\\scripts\\sudreg\\",
                "/opt/oriphiel-ai/scripts/sudreg/",
            ],
            [
                "Data (JSON/lock/log)",
                "C:\\GIT_PROJEKTI\\oriphiel-platform\\data\\sudreg\\",
                "/opt/oriphiel-ai/data/sudreg/",
            ],
            ["Progress", "data\\sudreg\\run\\progress.json", ".../run/progress.json"],
            ["Lock", "data\\sudreg\\run\\lock.json", ".../run/lock.json"],
            ["Logovi", "data\\sudreg\\logs\\", ".../logs/"],
            ["Shema SQL", "scripts\\sudreg\\sql\\sudreg-schema.sql", ".../sql/..."],
            ["Postgres", "SSH → docker na VPS", "lokalni docker (-Local)"],
            ["Baza", "sudreg", "sudreg (odvojeno od oriphiel)"],
        ],
    )

    # 4 Scripts
    h("4. Pregled skripti", 1)
    add_table(
        ["Skripta", "Uloga"],
        [
            ["Update-Sudreg.ps1", "Glavni tok: bootstrap + dnevna delta"],
            ["Run-SudregDaily.ps1", "Dnevni omotač + log"],
            ["Sudreg-Control.ps1", "Status / Stop / Wipe / ClearStaleLock"],
            ["Check-Sudreg.ps1", "Provjera napretka + baza + data dir (detaljno)"],
            ["SudregRun.ps1", "Lock, progress, abort"],
            ["SudregPg.ps1", "Postgres (SSH ili -Local docker)"],
            ["Sync-SudregToPostgres.ps1", "Scrape subjekata → UPSERT"],
            ["setup-sudreg-db.ps1", "CREATE DATABASE + shema"],
            ["Get-SudregToken.ps1", "OAuth token (cache)"],
            ["Get-SudregSnapshots.ps1", "Lista snapshotova"],
            ["Get-SudregPromjene.ps1", "/promjene za jedan snapshot"],
            ["Get-SudregSubject.ps1", "Scrape jedne firme (MBS)"],
            ["Sync-SudregPromjeneAll.ps1", "Sve /promjene → JSON"],
            ["Compare-SudregPromjene.ps1", "Diff MBS (added/removed)"],
            ["Deploy-SudregToVps.ps1", "Upload s Windowsa"],
            ["vps-install-sudreg.sh", "Install na Linuxu (pwsh, cron)"],
        ],
    )
    p(
        "Ne koristi -MaxMbs za pravi sync — redoslijed MBS nije po važnosti; "
        "rezanje propušta prave promjene.",
        bold=True,
    )

    # 5 Control
    h("5. Kontrola (status / stop / wipe)", 1)
    p("Na VPS-u uvijek dodaj -Local (ili export SUDREG_LOCAL=1).")

    h("5.1 Windows", 2)
    code_block(
        r"""cd C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg

# Dokle je došlo, je li završilo, brojevi u bazi
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -Status

# Isti izlaz kao JSON
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -Status -AsJson

# Zahtjev za prekid (završi trenutni batch, pa stane)
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -Stop

# Odmah ubij PID
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -Stop -ForceKill

# Potpuno obriši bazu sudreg + ponovo kreiraj praznu shemu
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -Wipe -Force

# + obriši i lokalne JSON datoteke (ostaju run/ i logs/)
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -Wipe -Force -AlsoClearJsonCache

# Ukloni mrtvi lock (PID više ne živi)
powershell -NoProfile -ExecutionPolicy Bypass -File .\Sudreg-Control.ps1 -ClearStaleLock"""
    )

    h("5.2 VPS (Linux)", 2)
    code_block(
        """export SUDREG_LOCAL=1
cd /opt/oriphiel-ai/scripts/sudreg

pwsh -NoProfile -File ./Sudreg-Control.ps1 -Status -Local
pwsh -NoProfile -File ./Sudreg-Control.ps1 -Stop -Local
pwsh -NoProfile -File ./Sudreg-Control.ps1 -Stop -ForceKill -Local
pwsh -NoProfile -File ./Sudreg-Control.ps1 -Wipe -Force -Local
pwsh -NoProfile -File ./Sudreg-Control.ps1 -ClearStaleLock -Local"""
    )

    h("5.3 Što Status pokazuje", 2)
    bullet("radi li unos (running) ili je završeno / prekinuto / greška")
    bullet("fazu: fetch_promjene, bootstrap_baseline, sync_scrape, …")
    bullet("done/total, %, ok / err, snapshot_id")
    bullet("broj firmi / osoba u bazi, last_imported_snapshot_id")

    h("5.4 Provjera do kud je došao unos (Check-Sudreg)", 2)
    p(
        "Check-Sudreg.ps1 je proširena provjera: status/napredak + progress.json + "
        "data dir + sync_state + snapshoti (+ opcionalno detaljna baza)."
    )
    p("Brzi status (dokle je došlo):")
    code_block(
        r"""# Windows
powershell -NoProfile -ExecutionPolicy Bypass -File `
  C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg\Sudreg-Control.ps1 -Status

# VPS
export SUDREG_LOCAL=1
pwsh -File /opt/oriphiel-ai/scripts/sudreg/Sudreg-Control.ps1 -Status -Local"""
    )
    p("Puna provjera (napredak + baza + datoteke):")
    code_block(
        r"""# Windows
powershell -NoProfile -ExecutionPolicy Bypass -File `
  C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg\Check-Sudreg.ps1

# VPS — preporučeno
pwsh -File /opt/oriphiel-ai/scripts/sudreg/Check-Sudreg.ps1 -Local

# Samo status (bez dodatnih SQL-ova)
pwsh -File /opt/oriphiel-ai/scripts/sudreg/Check-Sudreg.ps1 -StatusOnly -Local

# Detaljna baza (status firmi, scrape ok/err, zadnje firme)
pwsh -File /opt/oriphiel-ai/scripts/sudreg/Check-Sudreg.ps1 -DbDetail -Local

# JSON izlaz (za skripte)
pwsh -File /opt/oriphiel-ai/scripts/sudreg/Check-Sudreg.ps1 -StatusOnly -AsJson -Local"""
    )
    p("Što gledati u izlazu:")
    bullet("STANJE: RADI / ZAVRSENO / PREKINUTO / NEUSPJEH / MIRUJE")
    bullet("phase: fetch_promjene | bootstrap_baseline | bootstrap_delta | sync_scrape | daily_delta …")
    bullet("done/total (%) — dokle je scrape stigao u trenutnoj fazi")
    bullet("ok / err — uspješni vs neuspješni scrapeovi")
    bullet("companies / people / activities / snapshots u bazi")
    bullet("last_imported_snapshot_id — zadnji uvezeni snapshot")
    bullet("progress.json + lock.json u data/sudreg/run/")
    bullet("zadnji logovi u data/sudreg/logs/")
    p("Progress datoteka (ručno):")
    code_block(
        "# VPS\n"
        "cat /opt/oriphiel-ai/data/sudreg/run/progress.json\n"
        "watch -n5 cat /opt/oriphiel-ai/data/sudreg/run/progress.json"
    )

    h("5.5 Korisni SQL SELECT-ovi (Sudreg)", 2)
    p("Datoteka: sql/useful-selects.sql — sync_state, snapshoti, firme, osobe, promjene, scrape greške.")
    code_block(
        "# VPS — svi SELECT-ovi odjednom\n"
        "pwsh -File /opt/oriphiel-ai/scripts/sudreg/Check-Sudreg.ps1 -RunUsefulSql -Local\n\n"
        "# ili direktno psql\n"
        "docker exec -i oriphiel-postgres psql -U oriphiel -d sudreg \\\n"
        "  -f - < /opt/oriphiel-ai/scripts/sudreg/sql/useful-selects.sql"
    )
    p("Najkorisniji SELECT-ovi:")
    code_block(
        "-- Dokle je sync\n"
        "SELECT key, value, updated_at FROM sync_state ORDER BY key;\n\n"
        "-- Brojevi\n"
        "SELECT count(*) FROM companies;\n"
        "SELECT count(*) FROM company_people;\n"
        "SELECT scrape_ok, count(*) FROM companies GROUP BY scrape_ok;\n\n"
        "-- Zadnje firme\n"
        "SELECT mbs, left(naziv,40), oib, status, updated_at\n"
        "FROM companies ORDER BY updated_at DESC LIMIT 20;\n\n"
        "-- Pretraga po nazivu\n"
        "SELECT mbs, naziv, oib FROM companies WHERE naziv ILIKE '%oriphiel%';\n\n"
        "-- Scrape greške\n"
        "SELECT mbs, left(naziv,30), scrape_error FROM companies\n"
        "WHERE scrape_ok IS FALSE ORDER BY updated_at DESC LIMIT 20;\n\n"
        "-- Snapshoti\n"
        "SELECT id, timestamp, imported_at FROM snapshots ORDER BY id DESC LIMIT 10;\n\n"
        "-- Osobe u firmi (zamijeni MBS)\n"
        "SELECT * FROM company_people WHERE mbs = '080000000';"
    )

    # 6 Sync
    h("6. Unos podataka", 1)

    h("6.1 Preporučeno (prvi put i dnevno)", 2)
    p("Windows (Postgres ide preko SSH na VPS):")
    code_block(
        r"""powershell -NoProfile -ExecutionPolicy Bypass -File `
  C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg\Update-Sudreg.ps1 `
  -FetchPromjeneFirst -SkipExistingPromjene"""
    )
    p("VPS (lokalni docker):")
    code_block(
        """export SUDREG_LOCAL=1
pwsh -NoProfile -File /opt/oriphiel-ai/scripts/sudreg/Update-Sudreg.ps1 `
  -FetchPromjeneFirst -SkipExistingPromjene -Local"""
    )
    p("Dnevni omotač (log u data/sudreg/logs/):")
    code_block(
        r"powershell -NoProfile -ExecutionPolicy Bypass -File "
        r"C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg\Run-SudregDaily.ps1"
    )

    h("6.2 Samo shema / baza (jednom)", 2)
    code_block(
        r"""# Windows
powershell -NoProfile -ExecutionPolicy Bypass -File .\setup-sudreg-db.ps1

# VPS
pwsh -File /opt/oriphiel-ai/scripts/sudreg/setup-sudreg-db.ps1 -Local"""
    )

    # 7 Deploy
    h("7. Deploy na VPS", 1)
    p("Cilj: /opt/oriphiel-ai na root@186.240.157.80. Zahtijeva SSH ključ s PC-a.")
    code_block(
        r"""# Upload + install (pwsh, shema, cron 06:15)
powershell -NoProfile -ExecutionPolicy Bypass -File `
  C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg\Deploy-SudregToVps.ps1 -RunInstall

# Pokreni bootstrap u pozadini na VPS-u
powershell -NoProfile -ExecutionPolicy Bypass -File `
  C:\GIT_PROJEKTI\oriphiel-platform\scripts\sudreg\Deploy-SudregToVps.ps1 -SkipUpload -RunBootstrap"""
    )
    p("Ručno na VPS-u nakon upload-a:")
    code_block(
        """bash /opt/oriphiel-ai/scripts/sudreg/vps-install-sudreg.sh
pwsh -File /opt/oriphiel-ai/scripts/sudreg/Sudreg-Control.ps1 -Status -Local"""
    )
    p("Cron (nakon installa): svaki dan 06:15 → run-sudreg-daily.sh")

    # 8 Scenarios
    h("8. Tipični scenariji", 1)
    add_table(
        ["Situacija", "Što napraviti"],
        [
            [
                "Pokrenuo unos — gdje sam?",
                "Check-Sudreg.ps1 -Local  ILI  Sudreg-Control.ps1 -Status -Local",
            ],
            [
                "Želim nasilno stati",
                "-Stop; ako ne stane ~2 min: -Stop -ForceKill",
            ],
            [
                "Želim ispočetka",
                "1) -Stop -ForceKill  2) -Wipe -Force  3) Update-Sudreg …",
            ],
            [
                "Drug unos kaže da već traje, a proces je mrtav",
                "Sudreg-Control.ps1 -ClearStaleLock",
            ],
            [
                "Status s Windowsa bez brojeva iz baze",
                "Nema SSH ključa — na VPS-u: Check-Sudreg.ps1 -Local",
            ],
            [
                "Detaljna baza (firme, scrape)",
                "Check-Sudreg.ps1 -DbDetail -Local",
            ],
        ],
    )

    # 9 Checklist
    h("9. Checklist", 1)
    bullet("Docker kontejner oriphiel-postgres radi na VPS-u")
    bullet("Skripte u /opt/oriphiel-ai/scripts/sudreg/")
    bullet("Data u /opt/oriphiel-ai/data/sudreg/")
    bullet("pwsh instaliran (vps-install-sudreg.sh)")
    bullet("Cron 06:15 postavljen")
    bullet("Prvi sync: Update-Sudreg -FetchPromjeneFirst -SkipExistingPromjene -Local")
    bullet("Praćenje: Check-Sudreg.ps1 -Local  (ili Sudreg-Control.ps1 -Status -Local)")
    bullet("Detalj: Check-Sudreg.ps1 -DbDetail -Local")

    # 10 Notes
    h("10. Napomene", 1)
    bullet("Baza sudreg ≠ baza oriphiel (mail/FB messaging)")
    bullet("OAuth: env SUDREG_CLIENT_ID / SUDREG_CLIENT_SECRET, inače defaulti u Get-SudregToken.ps1")
    bullet("Na VPS-u uvijek -Local ili SUDREG_LOCAL=1")
    bullet("S Windowsa bez SSH ključa Status možda neće dohvatiti brojeve iz baze")
    bullet("Ovo je isti VPS kao Oriphiel Messaging (186.240.157.80), ali druga putanja (/opt/… vs /root/…)")

    # 11 vs messaging
    h("11. Sudreg vs Oriphiel Messaging (isti VPS)", 1)
    add_table(
        ["", "Sudreg", "Messaging"],
        [
            ["Putanja skripti", "/opt/oriphiel-ai/scripts/sudreg/", "/root/oriphiel-ai/oriphiel_messaging/"],
            ["Baza", "sudreg", "oriphiel"],
            ["Izvor", "Sudreg API", "IMAP / n8n"],
            ["Kontrola", "Sudreg-Control.ps1", "status JSON + bash"],
            ["Cron", "06:15 dnevno", "n8n IMAP trigger (live)"],
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        footer.add_run("\n— Kraj dokumenta — Sudreg —"),
        size=9,
        color=(120, 120, 120),
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
